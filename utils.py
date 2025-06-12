import os
import re
import fitz
import pandas as pd
from faker import Faker
import spacy
import ocrmypdf
from multiprocessing import Process, Queue
import hashlib
from datetime import datetime
import cv2
from pdf2image import convert_from_path
from ultralytics import YOLO
import numpy as np
from PIL import Image
import subprocess

import tempfile
from docx import Document
from docx.shared import Inches

from docx2pdf import convert


# === Configuration ===
fake = Faker("fr_FR")
DOSSIER_ANONYMISÉ = "fichiers_anonymises"
os.makedirs(DOSSIER_ANONYMISÉ, exist_ok=True)

# === Chargement modèle spaCy personnalisé ===
MODELE_PATH = os.path.join(os.path.dirname(__file__), "models", "model-best")
nlp = spacy.load(MODELE_PATH)

MODELE_PATH2 = os.path.join(os.path.dirname(__file__), "models", "model-best2")
nlp2 = spacy.load(MODELE_PATH2)

MODELE_PATH3 = os.path.join(os.path.dirname(__file__), "models", "runs1", "train","signature-detector","weights", "best.pt")
yolo = YOLO(MODELE_PATH3)


# === FEC ===
# --- Compteurs pour générer des identifiants anonymes ---
compteur_personne = 1
compteur_client = 1

# --- Fonctions anonymisation ---
def anonymiser_compte(val):
    if pd.notna(val):
        val_str = str(val)
        visible = val_str[:4]
        return visible + "X" * (len(val_str) - 4)
    return val

def anonymiser_piece(val):
    return "REF-" + str(fake.random_int(min=10000, max=99999)) if pd.notna(val) else val

def anonymiser_nom_generique(prefixe):
    global compteur_personne
    identifiant = f"{prefixe}{str(compteur_personne).zfill(3)}"
    compteur_personne += 1
    return identifiant

def anonymiser_client_generique():
    global compteur_client
    identifiant = f"Client{str(compteur_client).zfill(3)}"
    compteur_client += 1
    return identifiant

def detecter_separateur(chemin_fichier):
    with open(chemin_fichier, 'r', encoding='utf-8', errors='ignore') as f:
        ligne = f.readline()
        return "|" if ligne.count("|") > ligne.count("\t") else "\t"

# --- Fonction principale ---
def anonymiser_fichier_fec(chemin_fichier):
    global compteur_personne, compteur_client
    compteur_personne = 1
    compteur_client = 1

    try:
        sep = detecter_separateur(chemin_fichier)

        # Tentative avec utf-8, sinon latin1
        try:
            df = pd.read_csv(chemin_fichier, sep=sep, dtype=str, encoding='utf-8')
        except UnicodeDecodeError:
            df = pd.read_csv(chemin_fichier, sep=sep, dtype=str, encoding='latin1')

        # Colonnes à anonymiser
        colonnes_requises = ["CompteNum", "CompteLib", "CompAuxNum", "CompAuxLib", "PieceRef", "EcritureLib"]
        for col in colonnes_requises:
            if col not in df.columns:
                print(f"⚠️ Colonne manquante : {col} ➤ créée vide.")
                df[col] = ""

        # Anonymisation
        df["CompteNum"] = df["CompteNum"].apply(anonymiser_compte)
        df["CompAuxNum"] = df["CompAuxNum"].apply(anonymiser_compte)
        df["CompteLib"] = df["CompteLib"].apply(lambda x: anonymiser_nom_generique("Client") if pd.notna(x) else x)
        df["CompAuxLib"] = df["CompAuxLib"].apply(lambda x: anonymiser_client_generique() if pd.notna(x) else x)
        df["PieceRef"] = df["PieceRef"].apply(anonymiser_piece)
        df["EcritureLib"] = df["EcritureLib"].apply(lambda x: "Libellé anonymisé" if pd.notna(x) else x)

        # Création dossier si nécessaire
        os.makedirs(DOSSIER_ANONYMISÉ, exist_ok=True)

        # Export anonymisé
        nom_fichier = os.path.basename(chemin_fichier)
        sortie = os.path.join(DOSSIER_ANONYMISÉ, f"anonymise_{nom_fichier}")
        df.to_csv(sortie, sep=sep, index=False, encoding='utf-8')
        print(f"✅ Fichier anonymisé enregistré dans : {sortie}")
        return sortie

    except Exception as e:
        print("❌ Erreur d'anonymisation FEC :", str(e))
        return None

# === PDF OCR ===


def ocr_worker(input_pdf, output_pdf, queue):
    try:
        ocrmypdf.ocr(
            input_pdf,
            output_pdf,
            language='fra',
            force_ocr=True,
            output_type='pdf',
            optimize=0,
            deskew=False,
            remove_background=False,
            jobs=1,
            skip_big=20.0,
            oversample=100
        )
        queue.put("ok")
    except Exception as e:
        print("❌ OCR worker failed :", e)
        queue.put("fail")

def anonymiser_pdf_ocr(chemin_pdf):
    try:
        PDF_OCR = chemin_pdf.replace(".pdf", "_OCR.pdf")
        PDF_SORTIE = os.path.join(DOSSIER_ANONYMISÉ, "anonymise_" + os.path.basename(chemin_pdf))

        queue = Queue()
        p = Process(target=ocr_worker, args=(chemin_pdf, PDF_OCR, queue))
        p.start()
        p.join(timeout=120)  # max 2 min OCR

        if not queue.empty() and queue.get() == "ok":
            print("✅ OCR terminé :", PDF_OCR)
        else:
            print("❌ OCR échoué ou timeout")
            return None



        # === Règles
        LABELS_SENSIBLES = {"NOM", "ADRESSE", "SIRET", "NSS", "DATE", "CODE_NAF", "URSSAF", "MATRICULE"}

        def est_montant(text):
            return bool(re.fullmatch(r"[\d\s.,-]+", text) and not re.fullmatch(r"\d{5,}", text)) or \
                re.search(r"(euros?|net|brut|montant|versé|payer|rémunération|salaire)", text.lower())

        def est_vrai_matricule(text):
            return bool(re.fullmatch(r"[A-Z]{2,}[0-9]{2,}", text.strip()))

        def est_vraie_adresse(text):
            return any(m in text.lower() for m in ["rue", "avenue", "boulevard", "chemin", "impasse", "allée", "place"])

        def est_info_non_sensible(text):
            return "xpert-ia" in text.lower() or "avenue magellan" in text.lower()

        # === Anonymisation du PDF
        doc = fitz.open(PDF_OCR)
        total_anonymise = 0

        for page_num, page in enumerate(doc):
            print(f"\n📄 Traitement page {page_num + 1}")
            blocks = page.get_text("dict")["blocks"]
            modifications = []
            x_offset, y_offset = -2, 8

            for block in blocks:
                if block["type"] != 0:
                    continue
                for line in block["lines"]:
                    line_text = " ".join(span["text"].strip() for span in line["spans"]).strip()

                    # Détection NOM
                    regex_nom = r"(Madame|Monsieur|M\.|Mme)\s+([A-Z][a-zéèêëàâäîïôöûüç'’\-]+\s+){0,3}[A-Z]{2,}(?:\s+[A-Z]{2,})*"
                    match_nom = re.search(regex_nom, line_text)
                    nom_detecte = match_nom.group() if match_nom else None
                    nom_remplace = False

                    # ✅ Correction ici : regex_adresse
                    regex_adresse = r"\b\d{1,4}\s+(rue|avenue|boulevard|chemin|impasse|allée|place)\s+[A-ZÉÈA-Za-zàâäéèêëïîôöùûüç'’\-]+"
                    match_adresse = re.search(regex_adresse, line_text, re.IGNORECASE)
                    adresse_detectee = match_adresse.group() if match_adresse else None
                    adresse_remplacee = False

                    for span in line["spans"]:
                        text = span["text"].strip()
                        x0, y0 = span["bbox"][:2]
                        font_size = span["size"]
                        texte_anonymise = text

                        # Détection NLP
                        try:
                            if text.strip():
                                doc_spacy = nlp(text)
                            else:
                                continue
                        except Exception as e:
                            print(f"⛔ Erreur NLP sur : {text[:50]} → {e}")
                            continue

                        for ent in doc_spacy.ents:
                            val = ent.text.strip()
                            label = ent.label_.upper()
                            if (
                                    est_info_non_sensible(val)
                                    or (label == "ADRESSE" and not est_vraie_adresse(val))
                                    or (label == "MATRICULE" and not est_vrai_matricule(val))
                                    or est_montant(val)
                            ):
                                continue
                            if label in LABELS_SENSIBLES and val in texte_anonymise:
                                texte_anonymise = texte_anonymise.replace(val, "*" * len(val))


                        # Si nom détecté, anonymiser ses spans
                        if nom_detecte and text in nom_detecte:
                            print(f"🔒 Partie du NOM détectée : {text}")
                            texte_anonymise = ""

                        # Si adresse détectée, anonymiser ses spans
                        if adresse_detectee and text in adresse_detectee:
                            print(f"🔒 Partie de l'ADRESSE détectée : {text}")
                            texte_anonymise = ""

                        # Appliquer la modification visuelle
                        if texte_anonymise != text:
                            page.add_redact_annot(span["bbox"], fill=(1, 1, 1))
                            modifications.append((x0 + x_offset, y0 + y_offset, texte_anonymise, font_size))
                            total_anonymise += 1

            page.apply_redactions()
            for x, y, txt, size in modifications:
                page.insert_text((x, y), txt, fontsize=size, color=(0, 0, 0))

        doc.save(PDF_SORTIE)
        if not os.path.exists(PDF_SORTIE) or os.path.getsize(PDF_SORTIE) < 10_000:
            print(f"❌ PDF OCR mal généré ou vide : {PDF_SORTIE}")
            return None
        else:
            print(f"✅ PDF OCR vérifié (taille = {os.path.getsize(PDF_SORTIE)} octets)")

        doc.close()
        print(f"\n✅ PDF anonymisé généré : {PDF_SORTIE} — Total : {total_anonymise} éléments remplacés.")
        return PDF_SORTIE


    except Exception as e:
        import traceback
        print("❌ Erreur PDF OCR :", e)
        traceback.print_exc()
        return None


# === PDF simple ===
def anonymiser_pdf_simple(chemin_pdf):
    try:
        PDF_SORTIE = os.path.join(DOSSIER_ANONYMISÉ, "anonymise_" + os.path.basename(chemin_pdf))
        doc = fitz.open(chemin_pdf)
        modifications = []
        x_offset, y_offset = -2, 8
        LABELS_SENSIBLES = ["NOM", "ADRESSE", "SIRET", "NSS", "DATE", "CODE_NAF", "URSSAF", "MATRICULE"]

        def est_montant(text):
            text = text.strip()
            return bool(
                re.fullmatch(r"[\d\s.,-]+", text) and not re.fullmatch(r"\d{5,}", text)
            ) or re.search(r"(euros?|net|brut|montant|versé|payer|rémunération|salaire)", text.lower())

        def est_vrai_matricule(text):
            return bool(re.fullmatch(r"[A-Z]{2,}[0-9]{2,}", text.strip()))

        def est_vraie_adresse(text):
            mots_adresse = ["rue", "avenue", "bd", "boulevard", "impasse", "chemin", "allée"]
            return any(mot in text.lower() for mot in mots_adresse) or bool(re.fullmatch(r"\d{5} [A-ZÉÈÀ\- ]+", text))

        def est_info_non_sensible(text):
            if "XPERT-IA" in text:
                return True
            if re.search(r"\d{1,3} avenue magellan", text.lower()):
                return True
            return False

        for page_number, page in enumerate(doc):
            print(f"\n📄 Traitement de la page {page_number + 1}")
            blocks = page.get_text("dict")["blocks"]
            for block in blocks:
                if block['type'] == 0:
                    for line in block['lines']:
                        for span in line['spans']:
                            text = span['text']
                            x0, y0 = span['bbox'][:2]
                            font_size = span['size']
                            doc_spacy = nlp(text)
                            texte_anonymise = text
                            for ent in doc_spacy.ents:
                                label, val = ent.label_, ent.text.strip()
                                if (
                                    est_info_non_sensible(val) or
                                    (label == "ADRESSE" and not est_vraie_adresse(val)) or
                                    (label == "MATRICULE" and not est_vrai_matricule(val)) or
                                    est_montant(val)
                                ):
                                    print(f"⛔ Ignoré : {val} ({label})")
                                    continue
                                if label in LABELS_SENSIBLES:
                                    texte_anonymise = texte_anonymise.replace(val, "*" * len(val))
                            if texte_anonymise != text:
                                print(f"🔒 Bloc anonymisé : {text.strip()} ➡️ {texte_anonymise.strip()}")
                                page.add_redact_annot(span['bbox'], fill=(1, 1, 1))
                                modifications.append((x0 + x_offset, y0 + y_offset, texte_anonymise, font_size))
            page.apply_redactions()
            for x0, y0, texte, font_size in modifications:
                page.insert_text((x0, y0), texte, fontsize=font_size, color=(0, 0, 0))

        doc.save(PDF_SORTIE)
        doc.close()
        print(f"\n✅ PDF anonymisé sauvegardé sous : {PDF_SORTIE}")
        return PDF_SORTIE

    except Exception as e:
        print("Erreur PDF simple :", str(e))
        return None
# === Détection type PDF ===
def anonymiser_pdf(chemin_pdf):
    try:
        with fitz.open(chemin_pdf) as doc:
            has_text = any(page.get_text().strip() for page in doc)

            if has_text:
                texte_page1 = doc[0].get_text().lower()
                print("📝 Texte page 1 =", texte_page1[:200])

                if "contrat" in texte_page1 and "travail" in texte_page1:
                    print("📄 Contrat détecté — traitement complet")
                    return anonymiser_contrat_complet(chemin_pdf, is_scanned=False)

                print("📄 PDF simple détecté — Anonymisation bulletin")
                return anonymiser_pdf_simple(chemin_pdf)

            else:
                # 📸 OCR sur la première page uniquement pour détection
                from pdf2image import convert_from_path
                import pytesseract

                images = convert_from_path(chemin_pdf, first_page=1, last_page=1, dpi=300)
                text_ocr = pytesseract.image_to_string(images[0], lang="fra").lower()
                print("📝 Texte OCR page 1 =", text_ocr[:200])

                if "contrat" in text_ocr and "travail" in text_ocr:
                    print("📄 Contrat scanné détecté — traitement complet")
                    return anonymiser_contrat_complet(chemin_pdf, is_scanned=True)

                print("📄 PDF scanné mais pas contrat — traitement bulletin OCR")
                return anonymiser_pdf_ocr(chemin_pdf)

    except Exception as e:
        print("❌ Erreur lors de la détection du type de PDF :", str(e))
        return None



def anonymiser_contrat_complet(chemin_pdf, is_scanned=True):
    print(f"🧾 Chemin reçu : {chemin_pdf}")
    print(f"🔍 is_scanned ? {is_scanned}")
    print(f"📤 OCR sortie : {chemin_pdf.replace('.pdf', '_OCR.pdf') if is_scanned else chemin_pdf}")
    print("📥 Fichier source existe ?", os.path.exists(chemin_pdf))

    try:
        PDF_OCR = chemin_pdf.replace(".pdf", "_OCR.pdf") if is_scanned else chemin_pdf
        PDF_TEMP = PDF_OCR.replace(".pdf", "_TEMP.pdf")
        PDF_FINAL = os.path.join(DOSSIER_ANONYMISÉ, "anonymise_" + os.path.basename(chemin_pdf))

        if is_scanned:
            print("🔁 Lancement de l'OCR même pour grandes pages...")
            try:
                ocrmypdf.ocr(
                    chemin_pdf,
                    PDF_OCR,
                    force_ocr=True,
                    language="fra",
                    use_threads=True,
                    optimize=0,
                    jobs=1,
                    deskew=True,
                    pdf_renderer="sandwich",
                )
            except Exception as e:
                print("❌ Erreur OCR PDF :", e)
                return None

        # === Règles d’anonymisation
        LABELS = {"NOM", "ADRESSE", "SIRET", "NSS", "DATE", "CODE_NAF", "ENTREPRISE", "MATRICULE", "URSSAF"}
        EXCLUSIONS = {
            "CONTRAT DE TRAVAIL", "A TEMPS COMPLET", "REMUNERATION", "ARTICLE",
            "NON-CONCURRENCE", "FONCTIONS", "ENGAGEMENT", "ABSENCES", "DEPLACEMENTS"
        }
        REGEX_NOM_MANUEL = re.compile(r'\b(Monsieur|Mme|M\.?|Madame)\s+[A-Z][A-Z\-]+(?:\s+[A-Z][A-Z\-]+)?\b', re.IGNORECASE)
        REGEX_ENTREPRISE_MANUEL = re.compile(r'\b(la société|la sarl|l\'entreprise|le groupe|la sas|la sa)\s+([A-Z&\s\.\'\-]+)', re.IGNORECASE)
        REGEX_MATRICULE_MANUEL = re.compile(r'\b\d{10,14}\b')

        def corriger_erreurs_ocr(texte):
            texte = re.sub(r"\s+", " ", texte)
            corrections = {
                "N °": "N°", "n °": "n°", "S I R E T": "SIRET", "0 ": "0",
                "S A R L": "SARL", "S . A . S": "SAS", "S . A": "SA", "N . S . S": "NSS",
            }
            for k, v in corrections.items():
                texte = texte.replace(k, v)
            return texte.strip()

        # === Étape 1 : Anonymisation texte via spaCy
        doc = fitz.open(PDF_OCR)
        for page in doc:
            blocks = page.get_text("dict")["blocks"]
            modifications = []
            for block in blocks:
                if block['type'] != 0:
                    continue
                for line in block['lines']:
                    spans = line['spans']
                    full_line = "".join([s['text'] for s in spans])
                    full_line = corriger_erreurs_ocr(full_line)
                    doc_spacy = nlp2(full_line)
                    texte_anonymise = ""
                    last_idx = 0
                    used_spans = []

                    for ent in doc_spacy.ents:
                        if ent.label_ in LABELS and not any(ex in ent.text.upper() for ex in EXCLUSIONS):
                            texte_anonymise += full_line[last_idx:ent.start_char] + "*******"
                            last_idx = ent.end_char
                            used_spans.append((ent.start_char, ent.end_char))
                    texte_anonymise += full_line[last_idx:]

                    for match in REGEX_NOM_MANUEL.finditer(full_line):
                        s, e = match.span()
                        if not any(us <= s < ue or us < e <= ue for us, ue in used_spans):
                            texte_anonymise = texte_anonymise.replace(match.group(), "*******")

                    for match in REGEX_ENTREPRISE_MANUEL.finditer(full_line):
                        texte_anonymise = texte_anonymise.replace(match.group(), f"{match.group(1)} *******")

                    for match in REGEX_MATRICULE_MANUEL.finditer(full_line):
                        s, e = match.span()
                        if not any(us <= s < ue or us < e <= ue for us, ue in used_spans):
                            texte_anonymise = texte_anonymise.replace(match.group(), "********")

                    if texte_anonymise != full_line:
                        for s in spans:
                            page.add_redact_annot(s['bbox'], fill=(1, 1, 1))
                        x0, y0 = spans[0]['bbox'][:2]
                        size = spans[0]['size']
                        modifications.append((x0, y0, texte_anonymise, size))

            page.apply_redactions()
            for x0, y0, texte, size in modifications:
                page.insert_text((x0, y0 + 8), texte, fontsize=size, color=(0, 0, 0))

        doc.save(PDF_TEMP)
        doc.close()

        # === Étape 2 : Masquage des signatures (YOLO)
        images = convert_from_path(PDF_TEMP, dpi=300)
        images_finales = []

        for i, image_pil in enumerate(images):
            print(f"\n📄 Traitement image page {i + 1}")
            img_cv = cv2.cvtColor(np.array(image_pil), cv2.COLOR_RGB2BGR)
            temp_img = f"temp_page_{i + 1}.jpg"
            cv2.imwrite(temp_img, img_cv)
            results = yolo.predict(temp_img, conf=0.25, save=False)[0]

            for box in results.boxes.xyxy.cpu().numpy():
                x1, y1, x2, y2 = map(int, box)
                cv2.rectangle(img_cv, (x1, y1), (x2, y2), (255, 255, 255), -1)
                cv2.putText(img_cv, "[signature masquee]", (x1, y2 - 5), cv2.FONT_HERSHEY_SIMPLEX, 0.4, (0, 0, 0), 1)

            img_pil_final = Image.fromarray(cv2.cvtColor(img_cv, cv2.COLOR_BGR2RGB))
            images_finales.append(img_pil_final)

        images_finales[0].save(PDF_FINAL, save_all=True, append_images=images_finales[1:])
        print(f"\n✅ Contrat anonymisé + signatures masquées : {PDF_FINAL}")
        return PDF_FINAL

    except Exception as e:
        print("❌ Erreur traitement contrat complet :", e)
        import traceback
        traceback.print_exc()
        return None


# === Utilitaires DSN ===
adresse_fields = [
    "S21.G00.06.003",  # Adresse de l'établissement
    "S21.G00.11.003",  # Adresse de la DADS-U
    "S21.G00.30.007",  # Ville du salarié
    "S21.G00.30.009",  # Code postal
    "S21.G00.30.017",  # Adresse complémentaire
    "S21.G00.85.003",  # Adresse établissement
    "S21.G00.85.005",  # Ville établissement
]


# === Fonction d’anonymisation des adresses
def anonymiser_adresses(texte):
    lignes = texte.splitlines()
    nouvelles_lignes = []
    for ligne in lignes:
        for code in adresse_fields:
            if ligne.startswith(f"{code},"):
                ligne = f"{code},'[ADRESSE]'"
        nouvelles_lignes.append(ligne)
    return "\n".join(nouvelles_lignes)


# === Utilitaires DSN ===
def hash_nir(nir):
    return hashlib.sha256(nir.encode()).hexdigest()[:12]

def tranche_age(date_str):
    try:
        jour, mois, annee = int(date_str[:2]), int(date_str[2:4]), int(date_str[4:8])
        age = datetime.now().year - annee
        if age < 25: return "'Moins de 25 ans'"
        elif age < 30: return "'25-29 ans'"
        elif age < 35: return "'30-34 ans'"
        elif age < 40: return "'35-39 ans'"
        elif age < 50: return "'40-49 ans'"
        else: return "'50 ans et plus'"
    except:
        return "'X'"

# === Codes sensibles DSN ===
codes_anonymisation_dsn = {
    "S10.G00.00.001": lambda val, i: "'ANON_APP'",
    "S10.G00.00.002": lambda val, i: "'ANON_APP'",
    "S10.G00.00.003": lambda val, i: "'XXXXXXXXX'",
    "S10.G00.01.001": lambda val, i: "'000000000'",
    "S10.G00.01.003": lambda val, i: "'ENTREPRISE_X'",
    "S10.G00.01.004": lambda val, i: "'X'",
    "S10.G00.01.005": lambda val, i: val,
    "S10.G00.01.006": lambda val, i: "'X'",
    "S10.G00.02.002": lambda val, i: "'DECLARANT_X'",
    "S10.G00.02.004": lambda val, i: "'dummy@email.com'",
    "S10.G00.02.005": lambda val, i: "'DUMMY_PHONE'",
    "S20.G00.05.004": lambda val, i: "'" + hash_nir(val.strip("'")) + "'",
    "S20.G00.07.001": lambda val, i: f"'SAL_{i:04d}'",
    "S20.G00.07.002": lambda val, i: "'DUMMY_PHONE'",
    "S20.G00.07.003": lambda val, i: "'dummy@email.com'",
    "S21.G00.30.001": lambda val, i: "'" + hash_nir(val.strip("'")) + "'",
    "S21.G00.30.002": lambda val, i: f"'SAL_{i:04d}'",
    "S21.G00.30.003": lambda val, i: f"'SAL_{i:04d}'",
    "S21.G00.30.004": lambda val, i: "'X'",
    "S21.G00.30.006": lambda val, i: tranche_age(val.strip("'")),
    "S21.G00.30.008": lambda val, i: "'X'",
    "S21.G00.30.010": lambda val, i: "'X'",
}

# === Regex génériques DSN ===
email_regex_dsn = re.compile(r"[\w\.-]+@[\w\.-]+")
phone_regex_dsn = re.compile(r"'\d{10}'")
siret_regex_dsn = re.compile(r"'\d{9}'|'\d{14}'")
date_naissance_regex_dsn = re.compile(r"'\d{8}'")


# === Fonction principale ===
def anonymiser_fichier_dsn(chemin_fichier):
    try:
        with open(chemin_fichier, 'r', encoding='utf-8') as f:
            lignes = f.readlines()

        lignes_anonymisees = []
        compteur_salarie = 1

        for ligne in lignes:
            ligne = ligne.strip()
            if not ligne or ',' not in ligne:
                lignes_anonymisees.append(ligne + "\n")
                continue

            code, val = map(str.strip, ligne.split(",", 1))
            valeur = val.strip()

            if code in codes_anonymisation_dsn:
                valeur_anonyme = codes_anonymisation_dsn[code](valeur, compteur_salarie)
                ligne_anonyme = f"{code},{valeur_anonyme}"
                if code == "S21.G00.30.001":
                    compteur_salarie += 1
            else:
                # Anonymisation générique
                if email_regex_dsn.search(valeur):
                    ligne_anonyme = f"{code},'dummy@email.com'"
                elif phone_regex_dsn.search(valeur):
                    ligne_anonyme = f"{code},'0000000000'"
                elif siret_regex_dsn.search(valeur):
                    ligne_anonyme = f"{code},'000000000'"
                elif date_naissance_regex_dsn.search(valeur):
                    ligne_anonyme = f"{code},'01011970'"
                else:
                    ligne_anonyme = ligne  # inchangé

            lignes_anonymisees.append(ligne_anonyme)

        # Enregistrement
        nom_fichier = os.path.basename(chemin_fichier)
        sortie = os.path.join(DOSSIER_ANONYMISÉ, f"anonymise_{nom_fichier}")
        # Appliquer anonymisation des adresses
        contenu_anonymise = "\n".join(lignes_anonymisees)
        contenu_anonymise = anonymiser_adresses(contenu_anonymise)
        lignes_anonymisees = contenu_anonymise.splitlines(keepends=True)

        with open(sortie, 'w', encoding='utf-8') as f_out:
            f_out.writelines(lignes_anonymisees)


        print(f"✅ Fichier DSN anonymisé : {sortie}")
        return sortie

    except Exception as e:
        print("Erreur d'anonymisation DSN :", str(e))
        return None


def anonymiser_Contrat(pdf_ocr_path, pdf_sortie_path):
    REGEX_NOM_MANUEL = re.compile(
        r'\b(Monsieur|M.|Madame|M\.?|Mr\.?)\s+(?:[A-Z][a-zéèêàîïç\-]+\s+)?[A-Z][A-Z\-]+\b',flags=re.IGNORECASE
    )
    REGEX_ENTREPRISE_MANUEL = re.compile(
        r'\b(la société|la sarl|l\'entreprise|le groupe|la sas|la sa)\s+([A-Z&\s\.\'\-]+)', flags=re.IGNORECASE
    )

    EXCLUSIONS = [
        "CONTRAT DE TRAVAIL", "A TEMPS COMPLET", "A DUREE INDETERMINEE", "CREATIONS DU SALARIE","DEPLACEMENTS",
        "ARTICLE", "REMUNERATION", "FONCTIONS", "Entre", "ABSENCES", "CONFIDENTIALITE", "NON-CONCURRENCE","NON - CONCURRENCE","CREATIONS DU SALARIE"
    ]
    LABELS_READABLE = {
        "NOM": "nom", "ADRESSE": "adresse", "SIRET": "siret", "NSS": "nss", "DATE": "date",
        "CODE_NAF": "code_naf", "ENTREPRISE": "entreprise", "MATRICULE": "matricule", "URSSAF": "urssaf"
    }

    doc = fitz.open(pdf_ocr_path)

    for page_number, page in enumerate(doc):
        print(f"\n📄 Traitement de la page {page_number + 1}")
        blocks = page.get_text("dict")["blocks"]
        modifications = []
        x_offset, y_offset = -2, 8

        for block in blocks:
            if block['type'] != 0:
                continue

            for line in block['lines']:
                spans = line['spans']
                full_line = "".join([s['text'] for s in spans])
                doc_spacy = nlp2(full_line)

                texte_anonymise = ""
                last_idx = 0
                used_spans = []

                for ent in doc_spacy.ents:
                    if ent.label_ in LABELS_READABLE and not any(ex in ent.text.upper() for ex in EXCLUSIONS):
                        texte_anonymise += full_line[last_idx:ent.start_char] + "*******"
                        last_idx = ent.end_char
                        used_spans.append((ent.start_char, ent.end_char))
                texte_anonymise += full_line[last_idx:]

                if not any(e.label_ == "NOM" for e in doc_spacy.ents):
                    for match in REGEX_NOM_MANUEL.finditer(full_line):
                        span_start, span_end = match.span()
                        if not any(s <= span_start < e or s < span_end <= e for s, e in used_spans):
                            texte_anonymise = texte_anonymise.replace(match.group(), "********")

                for match in REGEX_ENTREPRISE_MANUEL.finditer(full_line):
                    texte_anonymise = texte_anonymise.replace(match.group(), f"{match.group(1)} *********")

                if texte_anonymise != full_line:
                    print(f"🔒 Bloc anonymisé : {full_line.strip()} ➡️ {texte_anonymise.strip()}")
                    for s in spans:
                        page.add_redact_annot(s['bbox'], fill=(1, 1, 1))
                    x0, y0 = spans[0]['bbox'][:2]
                    font_size = spans[0]['size']
                    modifications.append((x0 + x_offset, y0 + y_offset, texte_anonymise, font_size))

        page.apply_redactions()
        for x0, y0, texte, font_size in modifications:
            page.insert_text((x0, y0), texte, fontsize=font_size, color=(0, 0, 0))

    doc.save(pdf_sortie_path)
    doc.close()
    print(f"\n✅ PDF anonymisé sauvegardé sous : {pdf_sortie_path}")


def anonymiser_word_docx(chemin_docx):
    """
    Anonymise un fichier Word (.docx uniquement), masque les signatures
    et retourne le chemin du fichier anonymisé.
    """
    try:
        if chemin_docx.lower().endswith(".doc") and not chemin_docx.lower().endswith(".docx"):
            print("❌ Format .doc non supporté. Veuillez convertir ce fichier en .docx.")
            return None

        # ---------- 1. Anonymisation texte ----------
        LABELS_SENSIBLES = {"NOM", "ADRESSE", "SIRET", "NSS", "DATE", "DATE_NAISSANCE", "CODE_NAF", "ENTREPRISE", "MATRICULE", "URSSAF"}
        PROTECTED_KEYWORDS = [
            "état récapitulatif", "état prévisionnel", "article L", "article R", "ASSURANCES", "objet",
            "PROCES-VERBAL DES DELIBERATIONS", "(DPE)", "Code civil", "Code du commerce", "diagnostic de performance énergétique",
            "L'ASSEMBLEE GENERALE ORDINAIRES", "expropriation", "grosse ou un exemplaire", "decret n° 87-713",
            "quote-part", "charges locatives", "PREMIERE RESOLUTION", "DEUXIEME RESOLUTION", "TROISIEME RESOLUTION",
            "QUATRIEME RESOLUTION", "APPROBATION DES COMPTES", "AFFECTATION DU RESULTAT", "CONVENTIONS REGLEMENTEES",
            "POUVOIR POUR FORMALITES"
        ]
        REGEX_NOM_MANUEL = re.compile(r'\b(Monsieur|M\.?|Madame|Mr\.?)\s+[A-Z][a-zéèêàîïç\-]+\s+[A-Z][A-Z\-]+\b')
        REGEX_PRENOM_NOM = re.compile(r"\b([A-Z][a-zéèêàîïç\-]+)\s+([A-Z]{2,}(?:\s+[A-Z]{2,})?)\b")

        doc = Document(chemin_docx)
        noms_detectes = set()

        for para in doc.paragraphs:
            texte = original_text = para.text
            if any(k.lower() in texte.lower() for k in PROTECTED_KEYWORDS):
                continue

            doc_spacy = nlp2(texte)
            new_text, offset, used_spans = texte, 0, []

            for ent in doc_spacy.ents:
                if ent.label_ in LABELS_SENSIBLES and ent.text.strip() not in ["Madame", "Monsieur"]:
                    start, end = ent.start_char + offset, ent.end_char + offset
                    new_text = new_text[:start] + "*" * len(ent.text) + new_text[end:]
                    offset += len("*" * len(ent.text)) - len(ent.text)
                    used_spans.append((ent.start_char, ent.end_char))
                    if ent.label_ == "NOM":
                        noms_detectes.add(ent.text)

            for match in REGEX_NOM_MANUEL.finditer(texte):
                matched = match.group().strip()
                if matched in new_text and not any(s <= match.start() < e or s < match.end() <= e for s, e in used_spans):
                    new_text = new_text.replace(matched, "*" * len(matched))

            for match in REGEX_PRENOM_NOM.finditer(texte):
                full_name = match.group().strip()
                if full_name not in noms_detectes and len(full_name) > 5:
                    new_text = new_text.replace(full_name, "*" * len(full_name))
                    noms_detectes.add(full_name)

            match_adresse_full = re.search(r"(Demeurant au\s+[^\n\.]+)", original_text)
            if match_adresse_full:
                adresse_expr = match_adresse_full.group(1).strip()
                if adresse_expr in new_text:
                    new_text = new_text.replace(adresse_expr, "*" * len(adresse_expr))
                else:
                    for fragment in adresse_expr.split(","):
                        fragment = fragment.strip()
                        if fragment and fragment in new_text:
                            new_text = new_text.replace(fragment, "*" * len(fragment))

            if new_text != original_text:
                print(f"🔒 Paragraphe modifié : {original_text.strip()} ➡️ {new_text.strip()}")
                para.text = new_text

        # Fichier docx anonymisé
        docx_anonyme = os.path.join(tempfile.gettempdir(), os.path.basename(chemin_docx).replace(".docx", "_anonyme.docx"))
        doc.save(docx_anonyme)

        # ---------- 2. Conversion en images pour traitement YOLO ----------
        pdf_temp = os.path.join(tempfile.gettempdir(), "temp_pdf.pdf")
        subprocess.run(
            ["libreoffice", "--headless", "--convert-to", "pdf", "--outdir", os.path.dirname(pdf_temp), docx_anonyme],
            check=True
        )

        images = convert_from_path(pdf_temp, dpi=300)
        yolo_detecte_signature = False
        images_finales = []

        for i, image_pil in enumerate(images):
            img_cv = cv2.cvtColor(np.array(image_pil), cv2.COLOR_RGB2BGR)
            temp_img_path = os.path.join(tempfile.gettempdir(), f"page_{i + 1}.jpg")
            cv2.imwrite(temp_img_path, img_cv)
            results = yolo.predict(temp_img_path, conf=0.25, save=False)[0]

            if len(results.boxes) > 0:
                yolo_detecte_signature = True
                for box in results.boxes.xyxy.cpu().numpy():
                    x1, y1, x2, y2 = map(int, box)
                    cv2.rectangle(img_cv, (x1, y1), (x2, y2), (255, 255, 255), -1)
                    cv2.putText(img_cv, "[signature masquée]", (x1, y2 - 5), cv2.FONT_HERSHEY_SIMPLEX, 0.5, (0, 0, 0), 1)

            images_finales.append(Image.fromarray(cv2.cvtColor(img_cv, cv2.COLOR_BGR2RGB)))

        # ---------- 3. Génération du fichier final ----------
        nom_final = os.path.basename(chemin_docx).replace(".docx", "_anonymise.docx")
        sortie_docx = os.path.join(DOSSIER_ANONYMISÉ, nom_final)

        if not yolo_detecte_signature:
            os.replace(docx_anonyme, sortie_docx)
        else:
            final_docx = Document()
            for i, image in enumerate(images_finales):
                temp_image_path = os.path.join(tempfile.gettempdir(), f"masked_page_{i + 1}.jpg")
                image.save(temp_image_path, "JPEG")
                final_docx.add_paragraph().add_run().add_picture(temp_image_path, width=Inches(6.5))
            final_docx.save(sortie_docx)

        print(f"✅ Fichier final enregistré : {sortie_docx}")
        return sortie_docx

    except Exception as e:
        print(f"❌ Erreur dans anonymiser_word_docx : {e}")
        import traceback; traceback.print_exc()
        return None


def anonymiser_fichier(chemin_fichier):
    ext = os.path.splitext(chemin_fichier)[1].lower()

    if ext in {".csv", ".txt"}:
        return anonymiser_fichier_fec(chemin_fichier)
    elif ext == ".pdf":
        return anonymiser_pdf(chemin_fichier)
    elif ext == ".edi":
        return anonymiser_fichier_dsn(chemin_fichier)
    elif ext in {".doc", ".docx"}:
        return anonymiser_word_docx(chemin_fichier)
    else:
        print("❓ Format non pris en charge :", ext)
        return None
